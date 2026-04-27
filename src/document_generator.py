#!/usr/bin/env python3
"""문서 생성 모듈 (Word/PDF)

JSON 보고서 데이터를 Word/PDF 문서로 변환
"""

import os
import re
import json
import argparse
import tempfile
import subprocess
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table


def _check_pandoc_available():
    """Pandoc 사용 가능 여부 확인"""
    try:
        import pypandoc
        return True
    except ImportError:
        return False
    except Exception as e:
        print(f"⚠️ pypandoc import 오류: {e}")
        return False


class DocumentGenerator:
    """Word/PDF 문서 생성기"""

    def __init__(self):
        self.first_heading_added = False
        # 이미지 기본 디렉토리 (프로젝트 루트 기준)
        self.image_base_dir = Path(__file__).parent.parent / "data"

    def _generate_report_title(self, question: str) -> str:
        """질문 기반으로 보고서 제목 생성

        Note:
            LLM이 제목을 생성하지 못했을 때를 대비한 백업(fallback) 메서드
            실제로는 LLM이 생성한 제목을 우선 사용하고, 없을 경우에만 호출됨

        Args:
            question: 사용자 질문

        Returns:
            보고서 제목
        """
        # 키워드 기반 제목 매핑
        question_lower = question.lower()

        if "최종" in question or "종합" in question or "전체" in question:
            return "AI/ML 프로젝트 최종 보고서"
        elif "주간" in question or "weekly" in question_lower:
            return "주간 업무 보고서"
        elif "월간" in question or "monthly" in question_lower:
            return "월간 업무 보고서"
        elif "임원" in question or "executive" in question_lower:
            return "임원 보고서"
        elif "cmb" in question_lower or "추천" in question:
            return "CMB 추천시스템 보고서"
        elif "테니스" in question or "모멘텀" in question:
            return "테니스 모멘텀 예측 프로젝트 보고서"
        elif "급이량" in question or "아쿠아" in question:
            return "급이량 분석 프로젝트 보고서"
        elif "rag" in question_lower or "챗봇" in question:
            return "RAG 시스템 구축 보고서"
        else:
            # 기본 제목
            return "프로젝트 보고서"

    def _should_include_image(self, image_info: Dict[str, Any], answer_text: str) -> bool:
        """이미지를 보고서에 포함할지 결정

        Args:
            image_info: 이미지 정보 딕셔너리
            answer_text: 답변 텍스트

        Returns:
            포함 여부 (True/False)
        """
        description = image_info.get('description', '').lower()
        source = image_info.get('source', '').lower()

        # 제외할 이미지 패턴
        exclude_keywords = [
            '일정', 'schedule', '계획표', '프로젝트 일정',
            '참석자', 'participant', '회의록',
            '목차', 'table of contents'
        ]

        # 설명이나 출처에 제외 키워드가 있으면 제외
        for keyword in exclude_keywords:
            if keyword in description or keyword in source:
                return False

        # 포함할 이미지 패턴 (성능, 결과, 분석 관련)
        include_keywords = [
            '결과', 'result', '성능', 'performance',
            '그래프', 'graph', '차트', 'chart',
            '분석', 'analysis', '시각화', 'visualization',
            '모델', 'model', '예측', 'prediction',
            'accuracy', 'precision', 'recall', 'f1',
            '분포', 'distribution', '비교', 'comparison'
        ]

        # 설명이나 출처에 포함 키워드가 있으면 포함
        for keyword in include_keywords:
            if keyword in description or keyword in source:
                return True

        # 기본적으로 제외
        return False

    def _shorten_image_caption(self, description: str, max_length: int = 100) -> str:
        """이미지 캡션을 짧게 요약

        Args:
            description: 원본 설명
            max_length: 최대 길이

        Returns:
            요약된 설명
        """
        if not description or len(description) <= max_length:
            return description

        # 첫 문장만 추출
        first_sentence = description.split('.')[0].split('。')[0]

        if len(first_sentence) <= max_length:
            return first_sentence

        # 그래도 길면 자르고 ... 추가
        return first_sentence[:max_length-3] + "..."

    def _add_image(self, doc: Document, image_path: str, description: str = None, max_width: float = 5.0):
        """문서에 이미지 추가

        Args:
            doc: Document 객체
            image_path: 이미지 상대 경로 (예: "notion_images/xxx.png")
            description: 이미지 설명 (캡션)
            max_width: 최대 너비 (인치)
        """
        # 절대 경로로 변환
        full_path = self.image_base_dir / image_path

        if not full_path.exists():
            print(f"⚠️ 이미지 파일을 찾을 수 없습니다: {full_path}")
            # 이미지가 없으면 설명만 표시
            if description:
                para = doc.add_paragraph()
                run = para.add_run(f"[이미지: {description}]")
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
            return

        try:
            # 이미지 추가
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(str(full_path), width=Inches(max_width))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 캡션 추가
            if description:
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(description)
                caption_run.font.size = Pt(9)
                caption_run.font.italic = True
                caption_run.font.color.rgb = RGBColor(100, 100, 100)

        except Exception as e:
            print(f"⚠️ 이미지 삽입 실패: {full_path}, 오류: {e}")
            # 실패 시 설명만 표시
            if description:
                para = doc.add_paragraph()
                run = para.add_run(f"[이미지: {description}]")
                run.font.size = Pt(10)
                run.font.italic = True

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """헤딩 추가"""
        heading = doc.add_heading(text, level=level)
        return heading

    def _add_paragraph(self, doc: Document, text: str, bold: bool = False, italic: bool = False):
        """단락 추가"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'NanumGothic'

        if bold:
            run.bold = True
        if italic:
            run.italic = True

        return para

    def _parse_markdown_table(self, table_text: str) -> List[List[str]]:
        """마크다운 테이블 파싱

        Args:
            table_text: 마크다운 테이블 텍스트

        Returns:
            2D 리스트 (행x열)
        """
        lines = table_text.strip().split('\n')
        rows = []

        for i, line in enumerate(lines):
            # 빈 줄 스킵
            if not line.strip():
                continue

            # 테이블 행이 아니면 스킵
            if '|' not in line:
                continue

            # 구분선 스킵 (예: |---|---|, | --- | --- |, |-----|-----|)
            stripped = line.strip()
            # 하이픈으로만 구성된 셀이 있으면 구분선으로 간주
            temp_cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if temp_cells and all(set(cell) <= set('-: ') for cell in temp_cells):
                continue

            # 셀 추출: | 기준으로 분리
            parts = line.split('|')

            # 맨 앞뒤 빈 문자열 제거 (|로 시작하고 끝나는 경우)
            if parts and not parts[0].strip():
                parts = parts[1:]
            if parts and not parts[-1].strip():
                parts = parts[:-1]

            # 각 셀의 앞뒤 공백만 제거 (빈 셀 유지)
            cells = [cell.strip() for cell in parts]

            if cells:
                rows.append(cells)

        return rows

    def _add_markdown_table(self, doc: Document, table_text: str):
        """마크다운 테이블을 Word 테이블로 변환

        Args:
            doc: Word 문서
            table_text: 마크다운 테이블 텍스트
        """
        rows_data = self._parse_markdown_table(table_text)

        if not rows_data or len(rows_data) == 0:
            print(f"⚠️ 테이블 파싱 실패 또는 빈 테이블")
            return

        # 모든 행의 열 개수 확인 (가장 많은 열을 기준으로)
        num_cols = max(len(row) for row in rows_data)
        num_rows = len(rows_data)

        # 각 행의 열 개수를 맞춤 (부족한 경우 빈 셀 추가)
        for row in rows_data:
            while len(row) < num_cols:
                row.append('')

        print(f"📊 테이블 생성: {num_rows}행 x {num_cols}열")

        # Word 테이블 생성
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # 데이터 채우기
        for i, row_data in enumerate(rows_data):
            for j, cell_data in enumerate(row_data):
                if j < num_cols and i < num_rows:
                    cell = table.rows[i].cells[j]
                    cell.text = str(cell_data) if cell_data else ''

                    # 첫 행은 헤더로 볼드 처리
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(10)
                                run.font.name = 'NanumGothic'
                    else:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                                run.font.name = 'NanumGothic'

        doc.add_paragraph()  # 테이블 뒤 간격

    def _add_formatted_text(self, doc: Document, text: str):
        """마크다운 형식이 포함된 텍스트를 Word에 추가

        Args:
            doc: Word 문서
            text: 마크다운 형식 텍스트
        """
        lines = text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # 테이블 시작 감지: | 로 시작하고 끝나는 행
            if line.strip().startswith('|') and '|' in line:
                # 테이블 블록 수집
                table_lines = []
                start_i = i

                # 연속된 테이블 행 수집
                while i < len(lines):
                    current_line = lines[i].strip()

                    # 리스트 마커가 있는 경우 제거
                    current_line = re.sub(r'^\s*[-*+]\s+', '', current_line)
                    current_line = re.sub(r'^\s+', '', current_line)

                    # 테이블 행인지 확인
                    if current_line.startswith('|') and '|' in current_line:
                        table_lines.append(current_line)
                        i += 1
                    else:
                        # 테이블이 끝남
                        break

                # 수집한 테이블 블록 처리
                if table_lines:
                    table_text = '\n'.join(table_lines)
                    print(f"🔍 테이블 감지 ({start_i}행부터 {len(table_lines)}줄):")
                    print(table_text[:200] + '...' if len(table_text) > 200 else table_text)
                    self._add_markdown_table(doc, table_text)
            else:
                # 일반 텍스트 처리 - 다음 테이블까지의 모든 행 수집
                text_lines = []
                while i < len(lines):
                    if lines[i].strip().startswith('|') and '|' in lines[i]:
                        # 다음 테이블 발견
                        break
                    text_lines.append(lines[i])
                    i += 1

                if text_lines:
                    paragraph_text = '\n'.join(text_lines)
                    self._add_formatted_paragraph(doc, paragraph_text)

    def _add_formatted_paragraph(self, doc: Document, text: str):
        """마크다운 형식(볼드, 이탤릭, 리스트)을 Word로 변환

        Args:
            doc: Word 문서
            text: 텍스트
        """
        lines = text.strip().split('\n')

        for line in lines:
            if not line.strip():
                continue

            # 수평선 감지 (---, ___, ***)
            if re.match(r'^[\-_*]{3,}\s*$', line):
                # 수평선은 스킵 (표시하지 않음)
                continue

            # 리스트 항목 감지
            list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
            if list_match:
                indent = len(list_match.group(1))
                content = list_match.group(3)
                para = doc.add_paragraph(style='List Bullet')
                self._add_inline_formatting(para, content)
                continue

            # 헤딩 감지
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                content = heading_match.group(2)
                heading = doc.add_heading(content, level=level)

                # 첫 번째 레벨 1 헤딩만 중앙 정렬
                if not self.first_heading_added and level == 1:
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.first_heading_added = True

                continue

            # 일반 단락
            para = doc.add_paragraph()
            self._add_inline_formatting(para, line)

    def _add_inline_formatting(self, paragraph, text: str):
        """인라인 마크다운 형식(볼드, 이탤릭) 처리

        Args:
            paragraph: Word 단락
            text: 텍스트
        """
        # 볼드+이탤릭: ***text*** 또는 ___text___
        # 볼드: **text** 또는 __text__
        # 이탤릭: *text* 또는 _text_

        pattern = r'(\*\*\*|___|__|\*\*|_|\*)(.*?)\1'
        last_pos = 0

        for match in re.finditer(pattern, text):
            # 일반 텍스트 추가
            if match.start() > last_pos:
                run = paragraph.add_run(text[last_pos:match.start()])
                run.font.size = Pt(11)
                run.font.name = 'NanumGothic'

            # 형식화된 텍스트 추가
            marker = match.group(1)
            content = match.group(2)

            run = paragraph.add_run(content)
            run.font.size = Pt(11)
            run.font.name = 'NanumGothic'

            if marker in ['***', '___']:
                run.bold = True
                run.italic = True
            elif marker in ['**', '__']:
                run.bold = True
            elif marker in ['*', '_']:
                run.italic = True

            last_pos = match.end()

        # 남은 텍스트 추가
        if last_pos < len(text):
            run = paragraph.add_run(text[last_pos:])
            run.font.size = Pt(11)
            run.font.name = 'NanumGothic'

    def generate_word_report(self, report_data: Dict[str, Any], output_path: str):
        """Word 보고서 생성

        Args:
            report_data: 보고서 데이터 (JSON)
            output_path: 출력 파일 경로
        """
        # Pandoc 사용 가능 여부 확인
        pandoc_available = _check_pandoc_available()
        print(f"🔍 PANDOC_AVAILABLE: {pandoc_available}")
        if pandoc_available:
            print("🔧 Pandoc 하이브리드 방식 시도 중...")
            try:
                self._generate_word_with_pandoc_and_tables(report_data, output_path)
                return
            except Exception as e:
                import traceback
                print(f"⚠️ Pandoc 변환 실패, 기본 방식으로 전환: {e}")
                print(f"상세 오류:\n{traceback.format_exc()}")
        else:
            print("⚠️ pypandoc을 사용할 수 없습니다.")

        # 기본 방식 (python-docx)
        print("🔧 python-docx 방식으로 Word 생성 (테이블 지원)")
        self._generate_word_basic(report_data, output_path)

    def _remove_first_heading(self, text: str) -> str:
        """불필요한 헤딩 제거 및 정리

        Args:
            text: 마크다운 텍스트

        Returns:
            수정된 텍스트
        """
        import re
        lines = text.split('\n')
        result_lines = []
        skip_mode = False  # Executive Summary 섹션 전체 스킵용

        for line in lines:
            stripped = line.strip()

            # "임원 보고서", "최종 보고서" 등 불필요한 헤딩 제거
            if re.match(r'^#{1,3}\s*(임원\s*보고서|최종\s*보고서|일주일\s*보고서|주간\s*보고서)', stripped):
                continue

            # "Executive Summary (핵심 요약)" 섹션 시작 - 다음 ## 헤딩까지 스킵
            if re.match(r'^#{1,3}\s*Executive Summary', stripped, re.IGNORECASE):
                skip_mode = True
                continue

            # skip_mode 중 다음 섹션 헤딩 발견 시 스킵 모드 종료
            if skip_mode:
                # ## 또는 ### 로 시작하는 다른 섹션이 나오면 스킵 모드 종료
                if re.match(r'^#{1,3}\s*\d+\.', stripped):
                    skip_mode = False
                    # 현재 라인 처리로 넘어감 (continue 하지 않음)
                else:
                    # Executive Summary 섹션 내용 스킵
                    continue

            # 섹션 제목을 ## (레벨 2)로 통일
            # "### 1. 결과" → "## 1. 결과"
            if re.match(r'^#{1,4}\s*1\.', stripped):
                # 숫자 다음 텍스트 추출
                section_text = re.sub(r'^#{1,4}\s*', '', stripped)
                result_lines.append(f'## {section_text}')
                continue

            # "### 2. 주요 현황" → "## 2. 주요 현황"
            if re.match(r'^#{1,4}\s*2\.', stripped):
                section_text = re.sub(r'^#{1,4}\s*', '', stripped)
                result_lines.append(f'## {section_text}')
                continue

            # "### 3. 핵심 이슈 및 리스크" → "## 3. 핵심 이슈 및 리스크"
            if re.match(r'^#{1,4}\s*3\.', stripped):
                section_text = re.sub(r'^#{1,4}\s*', '', stripped)
                result_lines.append(f'## {section_text}')
                continue

            # "### 4. 추가 확인 필요 사항" → "## 4. 추가 확인 필요 사항"
            if re.match(r'^#{1,4}\s*4\.', stripped):
                section_text = re.sub(r'^#{1,4}\s*', '', stripped)
                result_lines.append(f'## {section_text}')
                continue

            result_lines.append(line)

        return '\n'.join(result_lines)

    def _fix_table_format(self, text: str) -> str:
        """유니코드 박스 문자를 마크다운 테이블로 변환하고 구분선 추가

        리스트 안의 테이블을 독립적인 테이블 블록으로 변환
        """
        # │ (유니코드 박스 문자)를 | (파이프)로 변환
        text = text.replace('│', '|')
        # ─ (유니코드 가로선)를 - (하이픈)로 변환
        text = text.replace('─', '-')

        lines = text.split('\n')
        fixed_lines = []
        in_table = False
        table_buffer = []

        for i, line in enumerate(lines):
            stripped = line.strip()

            # 테이블 행인지 확인 (리스트 마커 포함)
            is_table_line = False
            clean_line = line

            # 리스트 마커로 시작하는 테이블 행
            if re.match(r'^\s*[-*+]\s+\|', line):
                is_table_line = True
                clean_line = re.sub(r'^\s*[-*+]\s+', '', line)
            # 들여쓰기된 테이블 행
            elif re.match(r'^\s+\|', line):
                is_table_line = True
                clean_line = stripped
            # 일반 테이블 행
            elif stripped.startswith('|') and stripped.endswith('|'):
                is_table_line = True
                clean_line = stripped

            if is_table_line:
                if not in_table:
                    in_table = True
                    # 테이블 시작 전에 빈 줄 추가
                    if fixed_lines and fixed_lines[-1].strip():
                        fixed_lines.append('')
                table_buffer.append(clean_line)
            else:
                # 테이블이 끝났으면 버퍼 처리
                if in_table:
                    # 테이블 정리 및 추가
                    self._finalize_table_buffer(table_buffer, fixed_lines)
                    table_buffer = []
                    in_table = False
                    # 테이블 뒤에 빈 줄 추가
                    fixed_lines.append('')

                fixed_lines.append(line)

        # 마지막 테이블 처리
        if in_table and table_buffer:
            if fixed_lines and fixed_lines[-1].strip():
                fixed_lines.append('')
            self._finalize_table_buffer(table_buffer, fixed_lines)
            fixed_lines.append('')

        return '\n'.join(fixed_lines)

    def _finalize_table_buffer(self, table_buffer: list, output_lines: list):
        """테이블 버퍼를 정리하고 출력 라인에 추가"""
        if not table_buffer:
            return

        # 첫 번째 행이 헤더
        header = table_buffer[0]

        # 두 번째 행이 구분선인지 확인
        has_separator = False
        separator_idx = -1
        data_start_idx = 1

        # 버퍼에서 구분선 찾기 (첫 몇 행에서만)
        for idx in range(1, min(3, len(table_buffer))):
            line = table_buffer[idx].strip()
            # 구분선 패턴: |-----|-----|  또는 | --- | --- |
            if re.match(r'^\|[\s\-:|]+\|$', line) and '-' in line:
                has_separator = True
                separator_idx = idx
                data_start_idx = idx + 1
                break

        # 헤더 추가
        output_lines.append(header)

        # 구분선 추가 (있으면 원본 사용, 없으면 생성)
        if has_separator:
            output_lines.append(table_buffer[separator_idx])
        else:
            num_cols = header.count('|') - 1
            separator = '|' + '|'.join(['---' for _ in range(num_cols)]) + '|'
            output_lines.append(separator)

        # 나머지 데이터 행 추가
        for row in table_buffer[data_start_idx:]:
            output_lines.append(row)

    def _normalize_list_indentation(self, text: str) -> str:
        """마크다운 리스트 들여쓰기를 pandoc이 인식할 수 있도록 정규화

        2칸 들여쓰기를 4칸으로 변환
        """
        lines = text.split('\n')
        result_lines = []

        for line in lines:
            # 빈 줄은 그대로
            if not line.strip():
                result_lines.append(line)
                continue

            # 들여쓰기 감지
            stripped = line.lstrip()
            indent = len(line) - len(stripped)

            # 리스트 항목인지 확인 (- * + 또는 숫자.)
            is_list_item = bool(re.match(r'^[-*+]\s', stripped) or re.match(r'^\d+\.\s', stripped))

            if is_list_item and indent > 0:
                # 모든 들여쓰기를 4칸 단위로 정규화
                # 1-4칸 -> 4칸 (레벨 1)
                # 5-8칸 -> 8칸 (레벨 2)
                # 9-12칸 -> 12칸 (레벨 3)
                level = (indent - 1) // 4 + 1
                normalized_indent = level * 4
                result_lines.append(' ' * normalized_indent + stripped)
            else:
                # 들여쓰기 없는 리스트 항목이거나 일반 텍스트는 그대로
                result_lines.append(line)

        return '\n'.join(result_lines)

    def _generate_word_with_pandoc_and_tables(self, report_data: Dict[str, Any], output_path: str):
        """Pandoc + python-docx 하이브리드 방식으로 Word 생성

        표는 python-docx로 직접 생성하고, 나머지는 pandoc으로 변환
        표는 원래 마크다운에 있던 위치에 정확히 배치
        """
        import pypandoc
        print("🔧 하이브리드 방식: 표는 python-docx, 나머지는 Pandoc")

        # 결과 수집 및 표 추출
        results = report_data.get('results', [])
        markdown_content = []
        all_tables = []  # 모든 표를 순서대로 저장

        # 보고서 제목 추가 (LLM이 생성한 제목 사용, 없으면 기본 제목)
        if results and results[0].get('success'):
            # LLM이 생성한 제목 우선 사용
            title = results[0].get('title')
            if not title:
                # 제목이 없으면 질문 기반 생성
                title = self._generate_report_title(results[0]['question'])

            author = report_data.get('author', 'Unknown')
            created_date = report_data.get('created_date', datetime.now().strftime("%Y-%m-%d"))

            # 날짜 필터 정보 추출
            date_filter = results[0].get('date_filter', None)

            # 제목, 작성자, 작성일, 날짜 필터를 포함한 헤더를 placeholder로 추가
            # (Pandoc 변환 후 python-docx로 직접 스타일 적용)
            markdown_content.append('[REPORT_HEADER]')
            markdown_content.append(f'TITLE:{title}')
            markdown_content.append(f'AUTHOR:{author}')
            markdown_content.append(f'DATE:{created_date}')
            if date_filter:
                markdown_content.append(f'DATEFILTER:{date_filter}')
            markdown_content.append('[/REPORT_HEADER]')
            markdown_content.append('\n')

        # 각 답변을 처리하며 표를 placeholder로 치환
        for result in results:
            if result.get('success'):
                answer = result.get('answer', 'N/A')
                # 테이블 형식 수정
                answer = self._fix_table_format(answer)
                # "임원보고서" 헤딩 제거
                answer = self._remove_first_heading(answer)
                # 리스트 들여쓰기 정규화 (2칸 -> 4칸)
                answer = self._normalize_list_indentation(answer)

                # 표를 placeholder로 치환
                answer_with_placeholders, tables = self._replace_tables_with_placeholders(answer, len(all_tables))
                all_tables.extend(tables)

                markdown_content.append(answer_with_placeholders)
                markdown_content.append('\n\n')

        full_markdown = '\n'.join(markdown_content)

        # 디버깅: 마크다운 일부 출력
        print("🔍 생성된 마크다운 샘플 (첫 1000자):")
        print(full_markdown[:1000])
        print("...")

        # 임시 마크다운 파일 생성
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp:
            tmp.write(full_markdown)
            tmp_path = tmp.name

        print(f"📝 임시 마크다운 파일: {tmp_path}")
        if all_tables:
            print(f"📊 총 {len(all_tables)}개의 표 발견")

        try:
            # Pandoc으로 변환
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)

            # Pandoc 변환 시 리스트 간격 옵션 추가
            extra_args = []

            pypandoc.convert_file(
                tmp_path,
                'docx',
                outputfile=str(output_file),
                extra_args=extra_args
            )

            print(f"📄 Pandoc 변환 완료, 표 삽입 중...")

            # Word 문서 열기
            doc = Document(output_path)

            # 헤더(제목, 작성자, 작성일) 스타일 적용
            self._format_report_header(doc)

            # Placeholder를 실제 표로 교체
            self._replace_placeholders_with_tables(doc, all_tables)

            # 리스트 스타일 간격 조정
            self._adjust_list_spacing(doc)

            # 수정된 문서 저장
            doc.save(output_path)
            print(f"✅ Word 문서 생성 완료: {output_path}")

        finally:
            # 임시 파일 삭제 (디버깅을 위해 일시적으로 보존)
            print(f"🔍 디버깅: 임시 마크다운 파일 보존됨 - {tmp_path}")
            # Path(tmp_path).unlink(missing_ok=True)

    def _replace_tables_with_placeholders(self, text: str, start_index: int = 0) -> tuple:
        """마크다운 텍스트에서 표를 placeholder로 치환

        Args:
            text: 마크다운 텍스트
            start_index: 표 번호 시작 인덱스

        Returns:
            (치환된 텍스트, 추출된 표 리스트)
        """
        lines = text.split('\n')
        result_lines = []
        tables = []
        i = 0
        table_index = start_index

        while i < len(lines):
            line = lines[i]

            # 테이블 시작 감지
            if line.strip().startswith('|') and '|' in line:
                table_lines = []

                # 연속된 테이블 행 수집
                while i < len(lines):
                    current_line = lines[i].strip()
                    if current_line.startswith('|') and '|' in current_line:
                        table_lines.append(current_line)
                        i += 1
                    else:
                        break

                if table_lines:
                    # 표를 저장하고 placeholder 삽입
                    table_text = '\n'.join(table_lines)
                    tables.append(table_text)
                    result_lines.append(f'[TABLE_{table_index}]')
                    table_index += 1
            else:
                result_lines.append(line)
                i += 1

        return '\n'.join(result_lines), tables

    def _extract_tables_from_markdown(self, text: str) -> List[str]:
        """마크다운 텍스트에서 테이블 추출"""
        tables = []
        lines = text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # 테이블 시작 감지
            if line.strip().startswith('|') and '|' in line:
                table_lines = []

                # 연속된 테이블 행 수집
                while i < len(lines):
                    current_line = lines[i].strip()
                    if current_line.startswith('|') and '|' in current_line:
                        table_lines.append(current_line)
                        i += 1
                    else:
                        break

                if table_lines:
                    tables.append('\n'.join(table_lines))
            else:
                i += 1

        return tables

    def _format_report_header(self, doc: Document):
        """보고서 헤더(제목, 작성자, 작성일) 포맷 적용

        Args:
            doc: Word 문서 객체
        """
        # [REPORT_HEADER] 블록 찾기 (Pandoc이 한 줄로 합칠 수 있음)
        header_para = None
        header_data = {}

        for para in doc.paragraphs:
            text = para.text.strip()

            # Pandoc이 한 줄로 합친 경우
            if '[REPORT_HEADER]' in text and '[/REPORT_HEADER]' in text:
                header_para = para
                # 정규식으로 추출
                title_match = re.search(r'TITLE:([^\s]+(?:\s+[^\s]+)*?)(?:\s+AUTHOR:|$)', text)
                author_match = re.search(r'AUTHOR:([^\s]+(?:\s+[^\s]+)*?)(?:\s+DATE:|$)', text)
                date_match = re.search(r'DATE:([^\s]+(?:\s+[^\s]+)*?)(?:\s+(?:DATEFILTER:|\[/REPORT_HEADER\])|$)', text)
                datefilter_match = re.search(r'DATEFILTER:([^\s]+(?:\s+[^\s]+)*?)(?:\s+\[/REPORT_HEADER\]|$)', text)

                if title_match:
                    header_data['title'] = title_match.group(1).strip()
                if author_match:
                    header_data['author'] = author_match.group(1).strip()
                if date_match:
                    header_data['date'] = date_match.group(1).strip()
                if datefilter_match:
                    header_data['datefilter'] = datefilter_match.group(1).strip()

                break

        # 헤더 데이터가 있으면 포맷 적용
        if header_data and header_para is not None:
            # 기존 헤더 단락의 위치에 새로운 포맷 삽입
            p_element = header_para._element
            parent = p_element.getparent()
            insert_index = parent.index(p_element)

            # 제목 단락 생성 (중앙 정렬)
            title_para = header_para
            title_para.clear()
            title_run = title_para.add_run(header_data.get('title', ''))
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title_run.font.name = 'NanumGothic'
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 작성자 단락 생성 (우측 정렬)
            author_para = doc.add_paragraph()
            author_text = f"작성자: {header_data.get('author', 'Unknown')}"
            author_run = author_para.add_run(author_text)
            author_run.font.size = Pt(10)
            author_run.font.name = 'NanumGothic'
            author_run.font.color.rgb = RGBColor(100, 100, 100)
            author_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # 작성일 단락 생성 (우측 정렬)
            date_para = doc.add_paragraph()
            date_text = f"작성일: {header_data.get('date', '')}"
            date_run = date_para.add_run(date_text)
            date_run.font.size = Pt(10)
            date_run.font.name = 'NanumGothic'
            date_run.font.color.rgb = RGBColor(100, 100, 100)
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # author_para를 title_para 바로 다음으로 이동
            author_element = author_para._element
            author_element.getparent().remove(author_element)
            parent.insert(insert_index + 1, author_element)

            # date_para를 author_para 바로 다음으로 이동
            date_element = date_para._element
            date_element.getparent().remove(date_element)
            parent.insert(insert_index + 2, date_element)

            next_insert_index = insert_index + 3

            # 날짜 필터가 있으면 추가 (우측 정렬)
            if header_data.get('datefilter'):
                datefilter_para = doc.add_paragraph()
                datefilter_text = f"수행 기간: {header_data.get('datefilter')}"
                datefilter_run = datefilter_para.add_run(datefilter_text)
                datefilter_run.font.size = Pt(10)
                datefilter_run.font.name = 'NanumGothic'
                datefilter_run.font.color.rgb = RGBColor(100, 100, 100)
                datefilter_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # datefilter_para를 date_para 바로 다음으로 이동
                datefilter_element = datefilter_para._element
                datefilter_element.getparent().remove(datefilter_element)
                parent.insert(next_insert_index, datefilter_element)
                next_insert_index += 1

            # 빈 줄 추가
            blank_para = doc.add_paragraph()
            blank_element = blank_para._element
            blank_element.getparent().remove(blank_element)
            parent.insert(next_insert_index, blank_element)

            print(f"📋 보고서 헤더 포맷 적용 완료")

    def _adjust_list_spacing(self, doc: Document):
        """Word 문서의 리스트 항목 간격 및 들여쓰기 조정

        Args:
            doc: Word 문서 객체
        """
        adjusted_count = 0
        for para in doc.paragraphs:
            # 리스트 스타일인지 확인 (다양한 스타일 이름 대응)
            if para.style:
                style_name = para.style.name
                # Pandoc은 "Compact", "Tight", "List" 등 다양한 스타일 사용
                if ('List' in style_name or
                    'Bullet' in style_name or
                    'Compact' in style_name or
                    'Tight' in style_name):
                    # 단락 전후 간격 추가
                    para.paragraph_format.space_before = Pt(4)
                    para.paragraph_format.space_after = Pt(4)
                    # 줄 간격 조정 (1.2 배수)
                    para.paragraph_format.line_spacing = 1.2

                    # 들여쓰기 간격 조정 (기본값의 절반으로 줄임)
                    if para.paragraph_format.left_indent:
                        # 현재 들여쓰기를 절반으로 줄임
                        para.paragraph_format.left_indent = Inches(para.paragraph_format.left_indent.inches * 0.5)

                    adjusted_count += 1

        if adjusted_count > 0:
            print(f"🔧 리스트 간격 및 들여쓰기 조정 완료: {adjusted_count}개 항목")
        else:
            print(f"ℹ️  리스트 스타일을 찾지 못했습니다. 모든 단락에 간격 추가를 시도합니다...")
            # 리스트 스타일이 없으면 모든 단락에 약간의 간격 추가
            for para in doc.paragraphs:
                if para.text.strip():  # 빈 단락 제외
                    # 기존 간격이 없거나 매우 좁으면 간격 추가
                    if para.paragraph_format.space_after is None or para.paragraph_format.space_after < Pt(3):
                        para.paragraph_format.space_after = Pt(2)

    def _replace_placeholders_with_tables(self, doc: Document, markdown_tables: List[str]):
        """Word 문서의 placeholder를 python-docx 표로 교체

        Args:
            doc: Word 문서 객체
            markdown_tables: 마크다운 표 리스트 (순서대로)
        """
        # Placeholder를 찾아서 교체 정보 수집
        placeholders_to_replace = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            # Placeholder 패턴 확인
            if text.startswith('[TABLE_') and text.endswith(']'):
                # 표 번호 추출
                try:
                    placeholder_num = int(text[7:-1])  # [TABLE_X]에서 X 추출
                except:
                    continue

                if placeholder_num >= len(markdown_tables):
                    print(f"⚠️ 표 {placeholder_num} 범위 초과 (전체 {len(markdown_tables)}개)")
                    continue

                placeholders_to_replace.append({
                    'paragraph': paragraph,
                    'table_num': placeholder_num,
                    'md_table': markdown_tables[placeholder_num]
                })

        # 수집한 placeholder를 표로 교체
        for placeholder_info in placeholders_to_replace:
            paragraph = placeholder_info['paragraph']
            md_table = placeholder_info['md_table']
            table_num = placeholder_info['table_num']

            # 테이블 파싱
            rows_data = self._parse_markdown_table(md_table)
            if not rows_data:
                print(f"⚠️ 표 {table_num} 파싱 실패")
                continue

            num_rows = len(rows_data)
            num_cols = max(len(row) for row in rows_data)

            # 각 행의 열 개수를 맞춤
            for row in rows_data:
                while len(row) < num_cols:
                    row.append('')

            # Placeholder 단락 위치에 테이블 삽입
            p_element = paragraph._element
            parent = p_element.getparent()

            # 새 테이블 생성
            tbl = doc.add_table(rows=num_rows, cols=num_cols)._element

            # Placeholder 단락 바로 앞에 테이블 삽입
            parent.insert(parent.index(p_element), tbl)

            # 테이블 객체 가져오기 및 스타일 설정
            new_table = Table(tbl, doc)

            try:
                new_table.style = 'Light Grid Accent 1'
            except:
                try:
                    new_table.style = 'Table Grid'
                except:
                    pass

            # 데이터 채우기
            for i, row_data in enumerate(rows_data):
                for j, cell_data in enumerate(row_data):
                    if j < num_cols and i < num_rows:
                        cell = new_table.rows[i].cells[j]
                        cell.text = str(cell_data) if cell_data else ''

                        # 첫 행은 헤더로 볼드 처리
                        if i == 0:
                            for cell_para in cell.paragraphs:
                                for run in cell_para.runs:
                                    run.bold = True
                                    run.font.size = Pt(10)
                                    run.font.name = 'NanumGothic'
                        else:
                            for cell_para in cell.paragraphs:
                                for run in cell_para.runs:
                                    run.font.size = Pt(9)
                                    run.font.name = 'NanumGothic'

            # Placeholder 단락 삭제
            p_element.getparent().remove(p_element)

            print(f"✅ 표 {table_num} 삽입 완료 ({num_rows}행 x {num_cols}열)")

    def _replace_tables_in_word(self, doc: Document, markdown_tables: List[str]):
        """Word 문서의 테이블을 python-docx로 재생성한 테이블로 교체"""
        # 텍스트로 렌더링된 테이블 행들을 찾아서 삭제하고 그 자리에 실제 테이블 삽입
        paragraphs_to_remove = []
        table_insert_positions = []

        i = 0
        while i < len(doc.paragraphs):
            paragraph = doc.paragraphs[i]
            text = paragraph.text.strip()

            # 테이블 시작 감지 (| 로 시작하는 줄)
            if text.startswith('|') and '|' in text and len(text) > 10:
                # 테이블 블록의 모든 단락 수집
                table_paragraphs = [paragraph]
                table_start_idx = i
                j = i + 1

                # 연속된 테이블 행들 찾기
                while j < len(doc.paragraphs):
                    next_para = doc.paragraphs[j]
                    next_text = next_para.text.strip()

                    if next_text.startswith('|') and '|' in next_text:
                        table_paragraphs.append(next_para)
                        j += 1
                    else:
                        break

                # 이 위치에 테이블 삽입 예정
                if table_paragraphs and markdown_tables:
                    table_insert_positions.append({
                        'start_para': table_paragraphs[0],
                        'paragraphs': table_paragraphs,
                        'markdown_table': markdown_tables.pop(0)
                    })

                i = j
            else:
                i += 1

        # 역순으로 처리 (인덱스 변경 방지)
        for pos_info in reversed(table_insert_positions):
            start_para = pos_info['start_para']
            paragraphs = pos_info['paragraphs']
            md_table = pos_info['markdown_table']

            # 테이블 파싱
            rows_data = self._parse_markdown_table(md_table)
            if not rows_data:
                continue

            num_rows = len(rows_data)
            num_cols = max(len(row) for row in rows_data)

            # 각 행의 열 개수를 맞춤
            for row in rows_data:
                while len(row) < num_cols:
                    row.append('')

            # 첫 번째 단락 위치에 테이블 삽입
            p_element = start_para._element
            parent = p_element.getparent()

            # 새 테이블 생성 (단락 앞에 삽입)
            tbl = doc.add_table(rows=num_rows, cols=num_cols)._element

            # 단락 바로 앞에 테이블 삽입
            parent.insert(parent.index(p_element), tbl)

            # 테이블 객체 가져오기
            new_table = Table(tbl, doc)

            try:
                new_table.style = 'Light Grid Accent 1'
            except:
                try:
                    new_table.style = 'Table Grid'
                except:
                    pass

            # 데이터 채우기
            for i, row_data in enumerate(rows_data):
                for j, cell_data in enumerate(row_data):
                    if j < num_cols and i < num_rows:
                        cell = new_table.rows[i].cells[j]
                        cell.text = str(cell_data) if cell_data else ''

                        # 첫 행은 헤더로 볼드 처리
                        if i == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.size = Pt(10)
                                    run.font.name = 'NanumGothic'
                        else:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                                    run.font.name = 'NanumGothic'

            # 테이블에 해당하는 텍스트 단락들 삭제
            for para in paragraphs:
                p = para._element
                p.getparent().remove(p)

    def _generate_word_with_pandoc(self, report_data: Dict[str, Any], output_path: str):
        """Pandoc을 사용한 Word 생성 (마크다운 완벽 지원)"""
        import pypandoc
        # 결과 수집
        results = report_data.get('results', [])
        markdown_content = []

        # 보고서 제목 추가 (질문 기반)
        if results and results[0].get('question'):
            title = self._generate_report_title(results[0]['question'])
            markdown_content.append(f'# {title}\n\n')

            # 작성자 및 작성일자 추가
            author = report_data.get('author', 'Unknown')
            created_date = report_data.get('created_date', datetime.now().strftime("%Y-%m-%d"))
            markdown_content.append(f'**작성자:** {author}  |  **작성일:** {created_date}\n\n')
            markdown_content.append('---\n\n')

        for result in results:
            if result.get('success'):
                answer = result.get('answer', 'N/A')
                # 테이블 형식 수정
                answer = self._fix_table_format(answer)
                # "임원보고서" 헤딩 제거 (첫 번째 # 헤딩 제거)
                answer = self._remove_first_heading(answer)
                markdown_content.append(answer)
                markdown_content.append('\n\n')  # 질문 사이 간격

        full_markdown = '\n'.join(markdown_content)

        # 임시 마크다운 파일 생성
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp:
            tmp.write(full_markdown)
            tmp_path = tmp.name

        # 디버깅: 마크다운 파일 경로 출력
        print(f"📝 임시 마크다운 파일: {tmp_path}")

        try:
            # Pandoc으로 변환
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)

            pypandoc.convert_file(
                tmp_path,
                'docx',
                outputfile=str(output_file),
                extra_args=['--reference-doc='] if False else []  # 필요시 템플릿 추가
            )

            print(f"📄 Word 문서 생성 완료 (Pandoc): {output_file}")

        finally:
            # 임시 파일 삭제 (디버깅 시 주석 처리)
            # Path(tmp_path).unlink(missing_ok=True)
            pass

    def _generate_word_basic(self, report_data: Dict[str, Any], output_path: str):
        """기본 방식으로 Word 생성 (python-docx)"""
        doc = Document()

        # 첫 번째 헤딩 플래그 초기화
        self.first_heading_added = False

        # 보고서 제목 추가 (LLM이 생성한 제목 우선 사용, 없으면 질문 기반)
        results = report_data.get('results', [])
        if results and results[0].get('success'):
            # LLM이 생성한 제목 우선 사용
            title = results[0].get('title')
            if not title:
                # 제목이 없으면 질문 기반 생성
                title = self._generate_report_title(results[0]['question'])

            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.first_heading_added = True

            # 작성자 및 작성일자 추가 (우측 정렬, 별도 줄)
            author = report_data.get('author', 'Unknown')
            created_date = report_data.get('created_date', datetime.now().strftime("%Y-%m-%d"))

            # 작성자 단락 (우측 정렬)
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            author_run = author_para.add_run(f"작성자: {author}")
            author_run.font.size = Pt(10)
            author_run.font.name = 'NanumGothic'
            author_run.font.color.rgb = RGBColor(100, 100, 100)

            # 작성일 단락 (우측 정렬)
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_para.add_run(f"작성일: {created_date}")
            date_run.font.size = Pt(10)
            date_run.font.name = 'NanumGothic'
            date_run.font.color.rgb = RGBColor(100, 100, 100)

            # 날짜 필터 정보 추가 (수행날짜)
            date_filter = results[0].get('date_filter', None)
            if date_filter:
                datefilter_para = doc.add_paragraph()
                datefilter_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                datefilter_run = datefilter_para.add_run(f"보고 기간: {date_filter}")
                datefilter_run.font.size = Pt(10)
                datefilter_run.font.name = 'NanumGothic'
                datefilter_run.font.color.rgb = RGBColor(100, 100, 100)

            # 제목 뒤 간격
            doc.add_paragraph()

        # 결과

        for i, result in enumerate(results, 1):
            if result.get('success'):
                # 답변만 표시 (질문과 검색된 문서 정보는 제외)
                answer = result.get('answer', 'N/A')
                # 테이블 형식 수정
                answer = self._fix_table_format(answer)

                # 마크다운 형식 처리
                self._add_formatted_text(doc, answer)

                # 이미지 첨부 (필터링 적용)
                images = result.get('images', [])
                if images:
                    # 관련성 있는 이미지만 필터링
                    relevant_images = [
                        img for img in images
                        if self._should_include_image(img, answer)
                    ]

                    if relevant_images:
                        doc.add_paragraph()  # 답변과 이미지 사이 간격

                        # 이미지 섹션 제목
                        para = doc.add_paragraph()
                        run = para.add_run("📊 핵심 그래프 및 결과")
                        run.font.size = Pt(12)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(70, 70, 70)

                        # 각 이미지 삽입
                        for img in relevant_images:
                            img_path = img.get('path')
                            img_desc = img.get('description')
                            img_source = img.get('source')

                            if img_path:
                                # 이미지 설명 짧게 요약
                                caption = self._shorten_image_caption(img_desc) if img_desc else "이미지"
                                if img_source:
                                    caption += f" (출처: {img_source})"

                                self._add_image(doc, img_path, caption, max_width=5.0)
                                doc.add_paragraph()  # 이미지 사이 간격

            else:
                # 오류 발생
                self._add_paragraph(doc, f"❌ 오류 발생: {result.get('error', 'Unknown error')}", bold=True)

            # 질문 사이 간격
            if i < len(results):
                doc.add_paragraph()
                doc.add_paragraph()

        # 저장
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_file)

        print(f"📄 Word 문서 생성 완료 (기본): {output_file}")

    def generate_pdf_report(self, report_data: Dict[str, Any], output_path: str):
        """PDF 보고서 생성 (Word를 먼저 만들고 PDF로 변환)

        Args:
            report_data: 보고서 데이터 (JSON)
            output_path: 출력 파일 경로
        """
        # 임시 Word 파일 생성
        temp_docx = output_path.replace('.pdf', '_temp.docx')
        self.generate_word_report(report_data, temp_docx)

        try:
            # Word를 PDF로 변환 (LibreOffice 사용)
            output_file = Path(output_path)
            output_dir = output_file.parent

            # LibreOffice로 변환
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_dir),
                temp_docx
            ]

            result = subprocess.run(cmd, capture_output=True, text=True)

            if result.returncode == 0:
                # 생성된 PDF 파일을 원하는 이름으로 변경
                generated_pdf = temp_docx.replace('.docx', '.pdf')
                if Path(generated_pdf).exists() and generated_pdf != output_path:
                    Path(generated_pdf).rename(output_path)

                print(f"📄 PDF 문서 생성 완료: {output_path}")
            else:
                print(f"⚠️ PDF 변환 실패: {result.stderr}")
                print(f"💡 Word 파일을 사용하세요: {temp_docx}")

        except Exception as e:
            print(f"⚠️ PDF 생성 실패: {e}")
            print(f"💡 Word 파일을 사용하세요: {temp_docx}")

        finally:
            # 임시 Word 파일 삭제 (옵션)
            # Path(temp_docx).unlink(missing_ok=True)
            pass


    def generate_from_markdown(
        self,
        markdown_text: str,
        output_path: str,
        title: str = "보고서",
        author: str = "Unknown",
        created_date: str = None,
    ):
        """마크다운 텍스트를 Word/PDF로 직접 변환

        Args:
            markdown_text: 마크다운 형식의 텍스트 (LLM 출력 등)
            output_path: 출력 파일 경로 (.docx 또는 .pdf)
            title: 문서 제목
            author: 작성자
            created_date: 작성일 (None이면 오늘)
        """
        if created_date is None:
            created_date = datetime.now().strftime("%Y-%m-%d")

        report_data = {
            "results": [
                {
                    "question_id": 1,
                    "question": title,
                    "title": title,
                    "answer": markdown_text,
                    "success": True,
                    "images": [],
                    "date_filter": None,
                }
            ],
            "author": author,
            "created_date": created_date,
        }

        if output_path.endswith(".pdf"):
            self.generate_pdf_report(report_data, output_path)
        else:
            self.generate_word_report(report_data, output_path)


def main():
    """CLI 진입점"""
    parser = argparse.ArgumentParser(description="보고서 문서 생성기")
    parser.add_argument("--json", type=str, required=True, help="입력 JSON 파일")
    parser.add_argument("--output", type=str, required=True, help="출력 파일 (.docx 또는 .pdf)")

    args = parser.parse_args()

    # JSON 로드
    with open(args.json, 'r', encoding='utf-8') as f:
        report_data = json.load(f)

    # 문서 생성
    generator = DocumentGenerator()

    if args.output.endswith('.pdf'):
        generator.generate_pdf_report(report_data, args.output)
    elif args.output.endswith('.docx'):
        generator.generate_word_report(report_data, args.output)
    else:
        print("❌ 지원되지 않는 파일 형식입니다. .docx 또는 .pdf를 사용하세요.")


if __name__ == "__main__":
    main()
