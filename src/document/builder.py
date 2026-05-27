"""python-docx Word 문서 빌더"""

import re
import tempfile
from pathlib import Path
from typing import List
from urllib.parse import urlparse

import requests
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table

from .markdown import parse_markdown_table


class WordDocBuilder:
    """python-docx 문서 빌딩 헬퍼. Document 인스턴스를 감싸고 관련 상태를 관리한다."""

    def __init__(self, doc: Document, image_base_dir: Path):
        self.doc = doc
        self.image_base_dir = image_base_dir
        self.first_heading_added = False

    # ── 기본 요소 추가 ────────────────────────────────────────────────────────

    def add_heading(self, text: str, level: int = 1):
        return self.doc.add_heading(text, level=level)

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'NanumGothic'
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return para

    def _resolve_image_path(self, image_path: str) -> Path | None:
        if not image_path:
            return None
        if image_path.startswith(("http://", "https://")):
            try:
                response = requests.get(image_path, timeout=15)
                response.raise_for_status()
                suffix = Path(urlparse(image_path).path).suffix or ".png"
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                tmp.write(response.content)
                tmp.close()
                return Path(tmp.name)
            except Exception as e:
                print(f"⚠️ 이미지 다운로드 실패: {image_path}, 오류: {e}")
                return None

        full_path = Path(image_path)
        if full_path.exists():
            return full_path
        return self.image_base_dir / image_path

    def add_image(self, image_path: str, description: str = None, max_width: float = 5.0):
        full_path = self._resolve_image_path(image_path)

        if full_path is None or not full_path.exists():
            if description:
                para = self.doc.add_paragraph()
                run = para.add_run(f"[이미지: {description}]")
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
            return

        try:
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(str(full_path), width=Inches(max_width))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if description:
                caption_para = self.doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(description)
                caption_run.font.size = Pt(9)
                caption_run.font.italic = True
                caption_run.font.color.rgb = RGBColor(100, 100, 100)
        except Exception as e:
            print(f"⚠️ 이미지 삽입 실패: {full_path}, 오류: {e}")
            if description:
                para = self.doc.add_paragraph()
                run = para.add_run(f"[이미지: {description}]")
                run.font.size = Pt(10)
                run.font.italic = True

    # ── 마크다운 렌더링 ───────────────────────────────────────────────────────

    def add_markdown_table(self, table_text: str):
        """마크다운 테이블 → Word 테이블"""
        rows_data = parse_markdown_table(table_text)
        if not rows_data:
            return

        num_cols = max(len(row) for row in rows_data)
        num_rows = len(rows_data)

        for row in rows_data:
            while len(row) < num_cols:
                row.append('')

        print(f"📊 테이블 생성: {num_rows}행 x {num_cols}열")
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        for i, row_data in enumerate(rows_data):
            for j, cell_data in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = str(cell_data) if cell_data else ''
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10 if i == 0 else 9)
                        run.font.name = 'NanumGothic'
                        if i == 0:
                            run.bold = True

        self.doc.add_paragraph()

    def add_inline_formatting(self, paragraph, text: str):
        """인라인 마크다운(볼드, 이탤릭) + HTML span → Word run"""
        combined = r'(<span\b[^>]*>.*?</span>)|(\*\*\*|___|__|\*\*|_|\*)(.*?)\2'
        last_pos = 0

        for match in re.finditer(combined, text, flags=re.DOTALL | re.IGNORECASE):
            if match.start() > last_pos:
                run = paragraph.add_run(text[last_pos:match.start()])
                run.font.size = Pt(11)
                run.font.name = 'NanumGothic'

            if match.group(1):
                span_tag = match.group(1)
                inner = re.sub(r'<[^>]+>', '', span_tag)
                style = re.search(r'style=["\']([^"\']*)["\']', span_tag)
                run = paragraph.add_run(inner)
                run.font.size = Pt(11)
                run.font.name = 'NanumGothic'
                if style:
                    style_str = style.group(1)
                    color_m = re.search(r'color\s*:\s*(#[0-9a-fA-F]{6})', style_str)
                    if color_m:
                        hex_color = color_m.group(1).lstrip('#')
                        run.font.color.rgb = RGBColor(
                            int(hex_color[0:2], 16),
                            int(hex_color[2:4], 16),
                            int(hex_color[4:6], 16),
                        )
                    if 'bold' in style_str:
                        run.bold = True
            else:
                marker = match.group(2)
                content = match.group(3)
                run = paragraph.add_run(content)
                run.font.size = Pt(11)
                run.font.name = 'NanumGothic'
                if marker in ['***', '___']:
                    run.bold = True
                    run.italic = True
                elif marker in ['**', '__']:
                    run.bold = True
                elif marker in ['*', '_']:
                    run.italic = True

            last_pos = match.end()

        if last_pos < len(text):
            run = paragraph.add_run(text[last_pos:])
            run.font.size = Pt(11)
            run.font.name = 'NanumGothic'

    def add_formatted_paragraph(self, text: str):
        """마크다운 형식(헤딩, 리스트, 볼드, 이탤릭) → Word 단락들"""
        for line in text.strip().split('\n'):
            if not line.strip():
                continue

            if re.match(r'^[\-_*]{3,}\s*$', line):
                continue

            list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
            if list_match:
                para = self.doc.add_paragraph(style='List Bullet')
                self.add_inline_formatting(para, list_match.group(3))
                continue

            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                content = heading_match.group(2)
                heading = self.doc.add_heading(content, level=level)
                if not self.first_heading_added and level == 1:
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.first_heading_added = True
                continue

            para = self.doc.add_paragraph()
            self.add_inline_formatting(para, line)

    def add_formatted_text(self, text: str):
        """테이블 감지 후 Word 테이블/단락으로 분기 렌더링"""
        lines = text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            if line.strip().startswith('|') and '|' in line:
                table_lines = []
                start_i = i

                while i < len(lines):
                    current_line = lines[i].strip()
                    current_line = re.sub(r'^\s*[-*+]\s+', '', current_line)
                    current_line = re.sub(r'^\s+', '', current_line)

                    if current_line.startswith('|') and '|' in current_line:
                        table_lines.append(current_line)
                        i += 1
                    else:
                        break

                if table_lines:
                    self.add_markdown_table('\n'.join(table_lines))
            else:
                text_lines = []
                while i < len(lines):
                    if lines[i].strip().startswith('|') and '|' in lines[i]:
                        break
                    text_lines.append(lines[i])
                    i += 1

                if text_lines:
                    self.add_formatted_paragraph('\n'.join(text_lines))

    # ── 문서 후처리 ───────────────────────────────────────────────────────────

    def format_report_header(self):
        """[REPORT_HEADER] placeholder를 제목/작성자/작성일 스타일로 교체"""
        header_para = None
        header_data = {}

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if '[REPORT_HEADER]' in text and '[/REPORT_HEADER]' in text:
                header_para = para
                title_match = re.search(r'TITLE:([^\s]+(?:\s+[^\s]+)*?)(?:\s+AUTHOR:|$)', text)
                author_match = re.search(r'AUTHOR:([^\s]+(?:\s+[^\s]+)*?)(?:\s+DATE:|$)', text)
                date_match = re.search(r'DATE:([^\s]+(?:\s+[^\s]+)*?)(?:\s+(?:DATEFILTER:|\[/REPORT_HEADER\])|$)', text)
                datefilter_match = re.search(r'DATEFILTER:([^\s]+(?:\s+[^\s]+)*?)(?:\s+\[/REPORT_HEADER\]|$)', text)

                if title_match:
                    header_data['title'] = title_match.group(1).strip()
                if author_match:
                    header_data['author'] = author_match.group(1).strip()
                if date_match:
                    header_data['date'] = date_match.group(1).strip()
                if datefilter_match:
                    header_data['datefilter'] = datefilter_match.group(1).strip()
                break

        if not (header_data and header_para is not None):
            return

        from docx.shared import Pt, RGBColor
        p_element = header_para._element
        parent = p_element.getparent()
        insert_index = parent.index(p_element)

        # 제목 (중앙 정렬)
        title_para = header_para
        title_para.clear()
        title_run = title_para.add_run(header_data.get('title', ''))
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.name = 'NanumGothic'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _insert_right_aligned(text: str, offset: int):
            para = self.doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(10)
            run.font.name = 'NanumGothic'
            run.font.color.rgb = RGBColor(100, 100, 100)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elem = para._element
            elem.getparent().remove(elem)
            parent.insert(insert_index + offset, elem)

        _insert_right_aligned(f"작성자: {header_data.get('author', 'Unknown')}", 1)
        _insert_right_aligned(f"작성일: {header_data.get('date', '')}", 2)

        next_offset = 3
        if header_data.get('datefilter'):
            _insert_right_aligned(f"수행 기간: {header_data['datefilter']}", next_offset)
            next_offset += 1

        blank_para = self.doc.add_paragraph()
        blank_elem = blank_para._element
        blank_elem.getparent().remove(blank_elem)
        parent.insert(insert_index + next_offset, blank_elem)

        print("📋 보고서 헤더 포맷 적용 완료")

    def adjust_list_spacing(self):
        """리스트 항목 간격 및 들여쓰기 조정"""
        adjusted_count = 0
        for para in self.doc.paragraphs:
            if para.style:
                style_name = para.style.name
                if any(k in style_name for k in ('List', 'Bullet', 'Compact', 'Tight')):
                    para.paragraph_format.space_before = Pt(4)
                    para.paragraph_format.space_after = Pt(4)
                    para.paragraph_format.line_spacing = 1.2
                    if para.paragraph_format.left_indent:
                        para.paragraph_format.left_indent = Inches(
                            para.paragraph_format.left_indent.inches * 0.5
                        )
                    adjusted_count += 1

        if adjusted_count == 0:
            for para in self.doc.paragraphs:
                if para.text.strip():
                    if para.paragraph_format.space_after is None or para.paragraph_format.space_after < Pt(3):
                        para.paragraph_format.space_after = Pt(2)

    def replace_placeholders_with_tables(self, markdown_tables: List[str]):
        """Word 문서의 [TABLE_N] placeholder를 python-docx 테이블로 교체"""
        placeholders_to_replace = []

        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if text.startswith('[TABLE_') and text.endswith(']'):
                try:
                    placeholder_num = int(text[7:-1])
                except ValueError:
                    continue

                if placeholder_num >= len(markdown_tables):
                    print(f"⚠️ 표 {placeholder_num} 범위 초과 (전체 {len(markdown_tables)}개)")
                    continue

                placeholders_to_replace.append({
                    'paragraph': paragraph,
                    'table_num': placeholder_num,
                    'md_table': markdown_tables[placeholder_num],
                })

        for info in placeholders_to_replace:
            paragraph = info['paragraph']
            md_table = info['md_table']
            table_num = info['table_num']

            rows_data = parse_markdown_table(md_table)
            if not rows_data:
                print(f"⚠️ 표 {table_num} 파싱 실패")
                continue

            num_rows = len(rows_data)
            num_cols = max(len(row) for row in rows_data)

            for row in rows_data:
                while len(row) < num_cols:
                    row.append('')

            p_element = paragraph._element
            parent = p_element.getparent()

            tbl = self.doc.add_table(rows=num_rows, cols=num_cols)._element
            parent.insert(parent.index(p_element), tbl)

            new_table = Table(tbl, self.doc)
            try:
                new_table.style = 'Light Grid Accent 1'
            except Exception:
                try:
                    new_table.style = 'Table Grid'
                except Exception:
                    pass

            for i, row_data in enumerate(rows_data):
                for j, cell_data in enumerate(row_data):
                    cell = new_table.rows[i].cells[j]
                    cell.text = str(cell_data) if cell_data else ''
                    for cell_para in cell.paragraphs:
                        for run in cell_para.runs:
                            run.font.size = Pt(10 if i == 0 else 9)
                            run.font.name = 'NanumGothic'
                            if i == 0:
                                run.bold = True

            p_element.getparent().remove(p_element)
            print(f"✅ 표 {table_num} 삽입 완료 ({num_rows}행 x {num_cols}열)")
