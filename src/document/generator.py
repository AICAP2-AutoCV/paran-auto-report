#!/usr/bin/env python3
"""문서 생성 오케스트레이터 (Word/PDF)"""

import json
import os
import re
import argparse
import tempfile
import subprocess
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

from .markdown import (
    extract_title_tag,
    fix_table_format,
    parse_markdown_table,
    normalize_list_indentation,
    remove_first_heading,
    replace_tables_with_placeholders,
    extract_tables_from_markdown,
)
from .builder import WordDocBuilder


def _check_pandoc_available() -> bool:
    try:
        import pypandoc
        return True
    except (ImportError, Exception):
        return False


class DocumentGenerator:
    """마크다운 보고서 → Word/PDF 변환 오케스트레이터"""

    def __init__(self):
        self.image_base_dir = Path(__file__).parent.parent.parent / "data"

    # ── 보조 헬퍼 ─────────────────────────────────────────────────────────────

    def _generate_report_title(self, question: str) -> str:
        """질문 키워드 기반 보고서 제목 생성 (LLM 제목이 없을 때 폴백용)"""
        q = question.lower()
        if "최종" in question or "종합" in question or "전체" in question:
            return "AI/ML 프로젝트 최종 보고서"
        elif "주간" in question or "weekly" in q:
            return "주간 업무 보고서"
        elif "월간" in question or "monthly" in q:
            return "월간 업무 보고서"
        elif "임원" in question or "executive" in q:
            return "임원 보고서"
        elif "cmb" in q or "추천" in question:
            return "CMB 추천시스템 보고서"
        elif "테니스" in question or "모멘텀" in question:
            return "테니스 모멘텀 예측 프로젝트 보고서"
        elif "급이량" in question or "아쿠아" in question:
            return "급이량 분석 프로젝트 보고서"
        elif "rag" in q or "챗봇" in question:
            return "RAG 시스템 구축 보고서"
        return "프로젝트 보고서"

    def _should_include_image(self, image_info: Dict[str, Any], answer_text: str) -> bool:
        description = image_info.get('description', '').lower()
        source = image_info.get('source', '').lower()

        exclude = ['일정', 'schedule', '계획표', '프로젝트 일정', '참석자', 'participant', '회의록', '목차', 'table of contents']
        if any(k in description or k in source for k in exclude):
            return False

        include = ['결과', 'result', '성능', 'performance', '그래프', 'graph', '차트', 'chart',
                   '분석', 'analysis', '시각화', 'visualization', '모델', 'model', '예측', 'prediction',
                   'accuracy', 'precision', 'recall', 'f1', '분포', 'distribution', '비교', 'comparison']
        return any(k in description or k in source for k in include)

    def _shorten_image_caption(self, description: str, max_length: int = 100) -> str:
        if not description or len(description) <= max_length:
            return description
        first_sentence = description.split('.')[0].split('。')[0]
        if len(first_sentence) <= max_length:
            return first_sentence
        return first_sentence[:max_length - 3] + "..."

    def _collect_result_images(self, report_data: Dict[str, Any], max_images: int = 4) -> List[Dict[str, Any]]:
        images = []
        seen = set()
        for result in report_data.get("results", []):
            for image in result.get("images", []) or []:
                key = image.get("path") or image.get("url") or image.get("source") or image.get("description")
                if not key or key in seen:
                    continue
                seen.add(key)
                images.append(image)
                if len(images) >= max_images:
                    return images
        return images

    def _add_activity_photo_table(self, doc, images: List[Dict[str, Any]]):
        """활동 사진을 원본 양식처럼 2열 그리드 별도 표로 추가."""
        if not images:
            return

        builder = WordDocBuilder(doc, self.image_base_dir)
        valid = []
        for img in images:
            ref = img.get("path") or img.get("url")
            caption = self._shorten_image_caption(
                img.get("description") or img.get("caption") or
                img.get("section_title") or img.get("page_title") or "", 60
            )
            full_path = builder._resolve_image_path(ref)
            if full_path and full_path.exists():
                valid.append((full_path, caption))

        if not valid:
            return

        photo_table = doc.add_table(rows=1, cols=2)
        photo_table.autofit = False
        self._style_all_table_cells(photo_table)
        self._set_col_widths_dxa(photo_table, [5384, 5103])

        header_row = photo_table.rows[0]
        self._set_row_height(header_row, 426)
        merged = photo_table.cell(0, 0).merge(photo_table.cell(0, 1))
        self._set_cell_text(merged, "활동 사진", bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9E3F0", size=11)

        col_widths_cm = [9.1, 8.6]

        for i in range(0, len(valid), 2):
            pair = valid[i:i + 2]

            img_row_idx = len(photo_table.rows)
            photo_table.add_row()
            self._set_row_height(photo_table.rows[img_row_idx], 3530)

            cap_row_idx = len(photo_table.rows)
            photo_table.add_row()
            self._set_row_height(photo_table.rows[cap_row_idx], 1117)

            for col, (full_path, caption) in enumerate(pair):
                img_cell = photo_table.cell(img_row_idx, col)
                self._set_cell_border(img_cell)
                self._set_cell_margin(img_cell)
                img_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                para = img_cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                try:
                    para.add_run().add_picture(str(full_path), width=Cm(col_widths_cm[col]))
                except Exception as e:
                    print(f"⚠️ 이미지 삽입 실패: {full_path}, 오류: {e}")
                    run = para.add_run(f"[이미지: {caption}]")
                    run.font.size = Pt(10)
                    run.font.italic = True

                self._set_cell_text(photo_table.cell(cap_row_idx, col), caption,
                                    size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    def _append_related_images(self, target, images: List[Dict[str, Any]], doc=None, max_width: float = 13.5):
        if not images:
            return

        builder = WordDocBuilder(doc or target, self.image_base_dir)
        for image in images:
            image_ref = image.get("path") or image.get("url")
            description = (
                image.get("description")
                or image.get("caption")
                or image.get("section_title")
                or image.get("page_title")
                or "이미지"
            )
            caption = self._shorten_image_caption(description, max_length=140)
            full_path = builder._resolve_image_path(image_ref)
            if full_path is None or not full_path.exists():
                para = target.add_paragraph()
                run = para.add_run(f"[이미지: {caption}]")
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
                continue

            try:
                image_para = target.add_paragraph()
                image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_para.paragraph_format.space_before = Pt(6)
                image_para.paragraph_format.space_after = Pt(2)
                run = image_para.add_run()
                run.add_picture(str(full_path), width=Cm(max_width))

                caption_para = target.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.paragraph_format.space_before = Pt(0)
                caption_para.paragraph_format.space_after = Pt(6)
                caption_run = caption_para.add_run(caption)
                caption_run.font.size = Pt(9)
                caption_run.font.italic = True
                caption_run.font.color.rgb = RGBColor(100, 100, 100)
            except Exception as e:
                print(f"⚠️ 이미지 삽입 실패: {full_path}, 오류: {e}")
                para = target.add_paragraph()
                run = para.add_run(f"[이미지: {caption}]")
                run.font.size = Pt(10)
                run.font.italic = True

    def _is_weekly_activity_report(self, report_data: Dict[str, Any]) -> bool:
        results = report_data.get('results', [])
        if not results or not results[0].get('success'):
            return False
        answer = results[0].get('answer', '')
        title = results[0].get('title', '')
        markers = ('주차 활동내용', '1. 주요활동', '가. 최초 계획', '나. 실제 활동내용')
        return any(marker in answer for marker in markers) or '주차별 활동 보고서' in title

    def _table_to_mapping(self, table_text: str) -> Dict[str, str]:
        rows = parse_markdown_table(table_text)
        mapping = {}
        for row in rows[1:]:
            if len(row) >= 2:
                mapping[row[0].strip()] = row[1].strip()
        return mapping

    def _table_value(self, rows: List[List[str]], row_key: str, default: str = "확인 필요") -> str:
        for row in rows[1:]:
            if row and row[0].strip() == row_key:
                return row[-1].strip() or default
        return default

    def _table_col_value(self, rows: List[List[str]], row_key: str, col_idx: int, default: str = "확인 필요") -> str:
        for row in rows[1:]:
            if row and row[0].strip() == row_key and len(row) > col_idx:
                return row[col_idx].strip() or default
        return default

    def _extract_heading_body(self, text: str, heading: str) -> str:
        pattern = rf'^#+\s*{re.escape(heading)}\s*$'
        lines = text.splitlines()
        start = None
        start_level = 0
        for idx, line in enumerate(lines):
            if re.match(pattern, line.strip()):
                start = idx + 1
                start_level = len(line.strip()) - len(line.strip().lstrip('#'))
                break
        if start is None:
            return ""
        body = []
        for line in lines[start:]:
            stripped = line.strip()
            if stripped.startswith('#'):
                level = len(stripped) - len(stripped.lstrip('#'))
                if level <= start_level:
                    break
            if stripped.startswith('|'):
                continue
            body.append(line)
        return '\n'.join(body).strip()

    def _infer_semester_and_week(
        self,
        answer: str,
        basic: Dict[str, str],
        created_date: str,
    ) -> tuple[str, str, str]:
        """보고 기간 기준으로 학기와 주차 라벨을 계산."""
        period = basic.get("보고 기간", "")
        date_match = re.search(r'(\d{4}-\d{2}-\d{2})', period or answer)
        base_date = None
        if date_match:
            try:
                base_date = datetime.strptime(date_match.group(1), "%Y-%m-%d")
            except ValueError:
                base_date = None
        if base_date is None:
            try:
                base_date = datetime.strptime(created_date[:10], "%Y-%m-%d")
            except ValueError:
                base_date = datetime.now()

        if base_date.month >= 9:
            semester_year = base_date.year
            semester_no = 2
            semester_start = datetime(base_date.year, 9, 1)
        elif base_date.month >= 3:
            semester_year = base_date.year
            semester_no = 1
            semester_start = datetime(base_date.year, 3, 1)
        else:
            semester_year = base_date.year - 1
            semester_no = 2
            semester_start = datetime(base_date.year - 1, 9, 1)

        week_match = re.search(r'(\d+)\s*주차', answer)
        if week_match:
            week_no = week_match.group(1)
        else:
            days = max((base_date.date() - semester_start.date()).days, 0)
            week_no = str(days // 7 + 1)

        return f"{semester_year}-{semester_no}", week_no, f"{week_no}주차 활동내용"

    def _set_cell_border(self, cell, color: str = "000000", size: str = "8"):
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            tag = f"w:{edge}"
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), size)
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), color)

    def _shade_cell(self, cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.first_child_found_in("w:shd")
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), fill)

    def _set_cell_text(
        self,
        cell,
        text: str,
        bold: bool = False,
        size: int = 10,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        fill: str = None,
    ):
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if fill:
            self._shade_cell(cell, fill)
        self._set_cell_border(cell)
        self._set_cell_margin(cell)
        lines = str(text or "").splitlines() or [""]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for idx, line in enumerate(lines):
            if idx:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = "NanumGothic"

    def _style_all_table_cells(self, table):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(cell)
                self._set_cell_margin(cell)

    def _set_cell_margin(self, cell, top=28, left=102, bottom=28, right=102):
        """셀 내부 여백 설정 (dxa 단위, 원본 28/102 기준)."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = tcPr.find(qn('w:tcMar'))
        if tcMar is None:
            tcMar = OxmlElement('w:tcMar')
            tcPr.append(tcMar)
        for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
            el = tcMar.find(qn(f'w:{side}'))
            if el is None:
                el = OxmlElement(f'w:{side}')
                tcMar.append(el)
            el.set(qn('w:w'), str(val))
            el.set(qn('w:type'), 'dxa')

    def _set_row_height(self, row, val_dxa):
        """행 최소 높이 설정."""
        trPr = row._tr.get_or_add_trPr()
        trH = trPr.find(qn('w:trHeight'))
        if trH is None:
            trH = OxmlElement('w:trHeight')
            trPr.append(trH)
        trH.set(qn('w:val'), str(val_dxa))
        trH.set(qn('w:hRule'), 'atLeast')

    def _set_para_spacing(self, para, before=0, after=0, line=240):
        """단락 여백 및 줄간격 설정."""
        pPr = para._p.get_or_add_pPr()
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing')
            pPr.append(sp)
        sp.set(qn('w:before'), str(before))
        sp.set(qn('w:after'), str(after))
        sp.set(qn('w:line'), str(line))
        sp.set(qn('w:lineRule'), 'auto')

    def _set_col_widths(self, table, widths_cm):
        """tblGrid 기준 열 너비 설정 (merged cells에서도 안전)."""
        tbl = table._tbl
        tblGrid = tbl.find(qn('w:tblGrid'))
        if tblGrid is None:
            tblGrid = OxmlElement('w:tblGrid')
            tbl.insert(0, tblGrid)
        for gc in list(tblGrid.findall(qn('w:gridCol'))):
            tblGrid.remove(gc)
        for w in widths_cm:
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(int(w * 566.9)))
            tblGrid.append(gc)
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:w'), str(int(sum(widths_cm) * 566.9)))
            tblW.set(qn('w:type'), 'dxa')

    def _set_col_widths_dxa(self, table, widths_dxa, indent_dxa=-102):
        """원본 양식의 dxa 기반 표 너비/들여쓰기 설정."""
        tbl = table._tbl
        tblGrid = tbl.find(qn('w:tblGrid'))
        if tblGrid is None:
            tblGrid = OxmlElement('w:tblGrid')
            tbl.insert(0, tblGrid)
        for gc in list(tblGrid.findall(qn('w:gridCol'))):
            tblGrid.remove(gc)
        for width in widths_dxa:
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(width))
            tblGrid.append(gc)

        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), str(sum(widths_dxa)))
        tblW.set(qn('w:type'), 'dxa')

        jc = tblPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            tblPr.append(jc)
        jc.set(qn('w:val'), 'left')

        ind = tblPr.find(qn('w:tblInd'))
        if ind is None:
            ind = OxmlElement('w:tblInd')
            tblPr.append(ind)
        ind.set(qn('w:w'), str(indent_dxa))
        ind.set(qn('w:type'), 'dxa')

    def _add_cell_paragraph(self, cell, text, bold=False, size=10.5, color_rgb=None):
        """셀 내부에 단락 추가 (여백 없음)."""
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "NanumGothic"
        run.font.size = Pt(size)
        if color_rgb:
            run.font.color.rgb = color_rgb
        return p

    def _add_tight_spacer(self, doc, size_pt=2):
        """표 사이 구조 분리용 얇은 빈 문단."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        pPr = para._p.get_or_add_pPr()
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing')
            pPr.append(sp)
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '0')
        sp.set(qn('w:line'), '40')
        sp.set(qn('w:lineRule'), 'exact')
        run = para.add_run("")
        run.font.size = Pt(size_pt)
        return para

    def _add_report_paragraphs(self, doc, text: str):
        cleaned = re.sub(r'^\s*[-*]\s+', '', text, flags=re.MULTILINE).strip()
        if not cleaned:
            return
        for block in re.split(r'\n\s*\n', cleaned):
            block = block.strip()
            if not block:
                continue
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.35
            run = para.add_run(block.replace('\n', ' '))
            run.font.name = "NanumGothic"
            run.font.size = Pt(10.5)

    def _generate_weekly_word_template(self, report_data: Dict[str, Any], output_path: str):
        """학교 제출용 주차별 활동 보고서 — 외부 표 1개 + 중첩 표 구조."""
        results = report_data.get('results', [])
        result = results[0]
        title = result.get('title') or "주차별 활동 보고서"
        answer = fix_table_format(result.get('answer', ''))
        extracted_title, answer = extract_title_tag(answer)
        title = extracted_title or title
        created_date = report_data.get('created_date', datetime.now().strftime("%Y-%m-%d"))
        author = report_data.get('author', 'Unknown')
        role = report_data.get('role')

        basic = {}
        plan_rows = []
        actual_rows = []
        for table_text in extract_tables_from_markdown(answer):
            rows = parse_markdown_table(table_text)
            if not rows:
                continue
            header = " ".join(rows[0])
            if "항목" in header and "내용" in header:
                basic = self._table_to_mapping(table_text)
            elif "계획" in header:
                plan_rows = rows
            elif "투입시간" in header or "목표달성" in header:
                actual_rows = rows

        challenge = basic.get("도전과제명") or title
        report_date = created_date
        semester, week_no, week_label = self._infer_semester_and_week(answer, basic, created_date)

        values = [
            report_data.get("team_name") or basic.get("팀명", "확인 필요"),
            report_data.get("department") or basic.get("학과", "확인 필요"),
            report_data.get("student_id") or basic.get("학번", "확인 필요"),
            (author if author != "Unknown" else None) or basic.get("성명", "확인 필요"),
        ]

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f"{semester} 파란학기제 주차별 보고서")
        title_run.bold = True
        title_run.font.name = "NanumGothic"
        title_run.font.size = Pt(18)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(f"제출일자 : {report_date}")
        date_run.bold = True
        date_run.font.name = "NanumGothic"
        date_run.font.size = Pt(10.5)

        # 메타 표 — 원본처럼 활동내용 표와 분리된 4행 표
        meta = doc.add_table(rows=4, cols=4)
        meta.autofit = False
        self._style_all_table_cells(meta)
        self._set_col_widths_dxa(meta, [3630, 2310, 2595, 1920])
        for row, height in zip(meta.rows, [341, 361, 398, 233]):
            self._set_row_height(row, height)

        for idx, h in enumerate(["팀명", "학과", "학번", "성명"]):
            self._set_cell_text(meta.cell(0, idx), h, bold=True,
                                align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9E3F0")
        for idx, v in enumerate(values):
            self._set_cell_text(meta.cell(1, idx), v, align=WD_ALIGN_PARAGRAPH.CENTER)

        merged = meta.cell(2, 0).merge(meta.cell(2, 3))
        self._set_cell_text(merged, "도전과제명", bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9E3F0")
        merged = meta.cell(3, 0).merge(meta.cell(3, 3))
        self._set_cell_text(merged, challenge, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

        self._add_tight_spacer(doc)

        # 활동내용 표 — 원본처럼 별도 1열 3행 표 안에 본문과 중첩 표를 배치
        activity = doc.add_table(rows=3, cols=1)
        activity.autofit = False
        self._style_all_table_cells(activity)
        self._set_col_widths_dxa(activity, [10466])
        for row, height in zip(activity.rows, [426, 8269, 17509]):
            self._set_row_height(row, height)

        self._set_cell_text(activity.cell(0, 0), week_label.replace(" ", ""), bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9E3F0", size=11)

        content = activity.cell(1, 0)
        content.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        p0 = content.paragraphs[0]
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run("1. 주요활동")
        r0.bold = True
        r0.font.name = "NanumGothic"
        r0.font.size = Pt(12)

        self._add_cell_paragraph(content, "가. 최초 계획 *보완계획서에 기재한 주차별 계획 내용 참고", size=10.5)

        # 중첩 계획 표
        plan_table = content.add_table(rows=2, cols=3)
        plan_table.autofit = False
        self._style_all_table_cells(plan_table)
        self._set_col_widths(plan_table, [1.5, 8.0, 8.1])
        for i, h in enumerate(["주차", "팀", "개인"]):
            self._set_cell_text(plan_table.cell(0, i), h, bold=True,
                                align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9D9D9")
        self._set_cell_text(plan_table.cell(1, 0), week_no, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_text(plan_table.cell(1, 1), self._table_value(plan_rows, "팀"), size=10)
        self._set_cell_text(plan_table.cell(1, 2), self._table_value(plan_rows, "개인"), size=10)

        self._add_cell_paragraph(content, "나. 실제 활동내용 및 목표달성 여부", size=10.5)

        # 중첩 실제활동 표 ("팀" 셀 세로 병합)
        actual_table = content.add_table(rows=3, cols=3)
        actual_table.autofit = False
        self._style_all_table_cells(actual_table)
        self._set_col_widths(actual_table, [8.0, 2.5, 7.1])

        team_cell = actual_table.cell(0, 0).merge(actual_table.cell(1, 0))
        self._set_cell_text(team_cell, "팀", bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9D9D9")
        merged_indiv = actual_table.cell(0, 1).merge(actual_table.cell(0, 2))
        self._set_cell_text(merged_indiv, "개인", bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9D9D9")
        self._set_cell_text(actual_table.cell(1, 1), "투입시간", bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9D9D9")
        self._set_cell_text(actual_table.cell(1, 2), "실제 활동내용 및 목표달성 여부", bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER, fill="D9D9D9")

        team_actual = self._table_col_value(actual_rows, "팀", 2)
        team_status = self._table_col_value(actual_rows, "팀", 3, default="")
        personal_actual = self._table_col_value(actual_rows, "개인", 2)
        hours = "확인 필요"
        status = "확인 필요"
        for row in actual_rows[1:]:
            if row and row[0].strip() == "개인":
                raw_hours = row[1].strip() if len(row) > 1 else ""
                if raw_hours and raw_hours not in ("-", "확인 필요"):
                    hours = raw_hours
                raw_status = row[3].strip() if len(row) > 3 else ""
                if raw_status and raw_status != "확인 필요":
                    status = raw_status
                break

        # 팀 목표달성 여부가 있고 개인 데이터가 없으면 팀 값으로 보완
        if status == "확인 필요" and team_status and team_status != "확인 필요":
            status = team_status

        team_text = team_actual
        self._set_cell_text(actual_table.cell(2, 0), team_text, size=10)
        self._set_cell_text(actual_table.cell(2, 1), hours, align=WD_ALIGN_PARAGRAPH.CENTER)
        personal_parts = [p for p in [personal_actual, status] if p and p != "확인 필요"]
        if personal_parts:
            personal_cell_text = "\n\n".join(personal_parts)
        elif role:
            personal_cell_text = f"[{role}] 활동 내용 확인 필요"
        else:
            personal_cell_text = "확인 필요"
        self._set_cell_text(actual_table.cell(2, 2),
                            personal_cell_text,
                            align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

        details = self._extract_heading_body(answer, "2. 세부내용")
        lessons = self._extract_heading_body(answer, "3. 배운점")
        detail_cell = activity.cell(2, 0)
        detail_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        if details:
            self._add_cell_paragraph(detail_cell, "2. 세부내용", bold=True, size=12,
                                     color_rgb=RGBColor(0xC0, 0x00, 0x00))
            self._add_report_paragraphs(detail_cell, details)

        if lessons:
            self._add_cell_paragraph(detail_cell, "")
            self._add_cell_paragraph(detail_cell, "3. 배운점", bold=True, size=12,
                                     color_rgb=RGBColor(0xC0, 0x00, 0x00))
            self._add_report_paragraphs(detail_cell, lessons)

        self._add_tight_spacer(doc)
        self._add_activity_photo_table(doc, self._collect_result_images(report_data))

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_file)
        print(f"📄 주차별 보고서 양식 Word 생성 완료: {output_file}")

    # ── Word 생성 전략 ─────────────────────────────────────────────────────────

    def generate_word_report(self, report_data: Dict[str, Any], output_path: str):
        """Pandoc 하이브리드 방식 우선, 실패 시 python-docx 기본 방식으로 폴백"""
        if report_data.get("use_template") or self._is_weekly_activity_report(report_data):
            print("🧾 주차별 활동 보고서 전용 양식으로 생성")
            self._generate_weekly_word_template(report_data, output_path)
            return

        if _check_pandoc_available():
            print("🔧 Pandoc 하이브리드 방식 시도 중...")
            try:
                self._generate_word_with_pandoc_and_tables(report_data, output_path)
                return
            except Exception as e:
                import traceback
                print(f"⚠️ Pandoc 변환 실패, 기본 방식으로 전환: {e}")
                print(traceback.format_exc())
        else:
            print("⚠️ pypandoc을 사용할 수 없습니다.")

        print("🔧 python-docx 방식으로 Word 생성")
        self._generate_word_basic(report_data, output_path)

    def _generate_word_with_pandoc_and_tables(self, report_data: Dict[str, Any], output_path: str):
        """Pandoc + python-docx 하이브리드: 본문은 Pandoc, 표는 python-docx로 삽입"""
        import pypandoc

        results = report_data.get('results', [])
        markdown_content = []
        all_tables = []

        if results and results[0].get('success'):
            title = results[0].get('title') or self._generate_report_title(results[0]['question'])
            author = report_data.get('author', 'Unknown')
            created_date = report_data.get('created_date', datetime.now().strftime("%Y-%m-%d"))
            date_filter = results[0].get('date_filter', None)

            markdown_content += [
                '[REPORT_HEADER]',
                f'TITLE:{title}',
                f'AUTHOR:{author}',
                f'DATE:{created_date}',
            ]
            if date_filter:
                markdown_content.append(f'DATEFILTER:{date_filter}')
            markdown_content += ['[/REPORT_HEADER]', '\n']

        for result in results:
            if result.get('success'):
                answer = result.get('answer', 'N/A')
                _, answer = extract_title_tag(answer)
                answer = fix_table_format(answer)
                answer = remove_first_heading(answer)
                answer = normalize_list_indentation(answer)

                answer_with_placeholders, tables = replace_tables_with_placeholders(answer, len(all_tables))
                all_tables.extend(tables)

                markdown_content.append(answer_with_placeholders)
                markdown_content.append('\n\n')

        full_markdown = '\n'.join(markdown_content)

        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp:
            tmp.write(full_markdown)
            tmp_path = tmp.name

        print(f"📝 임시 마크다운 파일: {tmp_path}")
        if all_tables:
            print(f"📊 총 {len(all_tables)}개의 표 발견")

        try:
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)

            pypandoc.convert_file(tmp_path, 'docx', outputfile=str(output_file))

            doc = Document(output_path)
            builder = WordDocBuilder(doc, self.image_base_dir)
            builder.format_report_header()
            builder.replace_placeholders_with_tables(all_tables)
            builder.adjust_list_spacing()
            doc.save(output_path)

            print(f"✅ Word 문서 생성 완료: {output_path}")
        finally:
            print(f"🔍 디버깅: 임시 마크다운 파일 보존됨 - {tmp_path}")

    def _generate_word_basic(self, report_data: Dict[str, Any], output_path: str):
        """python-docx만으로 Word 생성 (Pandoc 폴백)"""
        doc = Document()
        builder = WordDocBuilder(doc, self.image_base_dir)

        results = report_data.get('results', [])
        if results and results[0].get('success'):
            title = results[0].get('title') or self._generate_report_title(results[0]['question'])

            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            builder.first_heading_added = True

            author = report_data.get('author', 'Unknown')
            created_date = report_data.get('created_date', datetime.now().strftime("%Y-%m-%d"))
            date_filter = results[0].get('date_filter', None)

            def _right_line(text: str):
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = para.add_run(text)
                run.font.size = Pt(10)
                run.font.name = 'NanumGothic'
                run.font.color.rgb = RGBColor(100, 100, 100)

            _right_line(f"작성자: {author}")
            _right_line(f"작성일: {created_date}")
            if date_filter:
                _right_line(f"보고 기간: {date_filter}")
            doc.add_paragraph()

        for i, result in enumerate(results, 1):
            if result.get('success'):
                answer = fix_table_format(result.get('answer', 'N/A'))
                _, answer = extract_title_tag(answer)
                builder.add_formatted_text(answer)

                images = result.get('images', [])
                relevant_images = [img for img in images if self._should_include_image(img, answer)]
                if relevant_images:
                    doc.add_paragraph()
                    para = doc.add_paragraph()
                    run = para.add_run("📊 핵심 그래프 및 결과")
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(70, 70, 70)

                    for img in relevant_images:
                        img_path = img.get('path') or img.get('url')
                        if img_path:
                            caption = self._shorten_image_caption(img.get('description')) if img.get('description') else "이미지"
                            if img.get('source'):
                                caption += f" (출처: {img['source']})"
                            builder.add_image(img_path, caption, max_width=5.0)
                            doc.add_paragraph()
            else:
                builder.add_paragraph(f"❌ 오류 발생: {result.get('error', 'Unknown error')}", bold=True)

            if i < len(results):
                doc.add_paragraph()
                doc.add_paragraph()

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_file)
        print(f"📄 Word 문서 생성 완료 (기본): {output_file}")

    # ── PDF 생성 ──────────────────────────────────────────────────────────────

    def generate_pdf_report(self, report_data: Dict[str, Any], output_path: str):
        """Word를 먼저 만들고 LibreOffice로 PDF 변환"""
        temp_docx = output_path.replace('.pdf', '_temp.docx')
        self.generate_word_report(report_data, temp_docx)

        try:
            output_dir = Path(output_path).parent
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', str(output_dir), temp_docx],
                capture_output=True, text=True,
            )
            if result.returncode == 0:
                generated_pdf = temp_docx.replace('.docx', '.pdf')
                if Path(generated_pdf).exists() and generated_pdf != output_path:
                    Path(generated_pdf).rename(output_path)
                print(f"📄 PDF 문서 생성 완료: {output_path}")
            else:
                print(f"⚠️ PDF 변환 실패: {result.stderr}")
                print(f"💡 Word 파일을 사용하세요: {temp_docx}")
        except Exception as e:
            print(f"⚠️ PDF 생성 실패: {e}")
            print(f"💡 Word 파일을 사용하세요: {temp_docx}")

    # ── 공개 API ──────────────────────────────────────────────────────────────

    def generate_from_markdown(
        self,
        markdown_text: str,
        output_path: str,
        title: str = "보고서",
        author: str = "Unknown",
        created_date: str = None,
        student_id: str = None,
        department: str = None,
        team_name: str = None,
        role: str = None,
        images: List[Dict[str, Any]] = None,
    ):
        """마크다운 텍스트를 Word/PDF 파일로 변환"""
        if created_date is None:
            created_date = datetime.now().strftime("%Y-%m-%d")
        extracted_title, markdown_text = extract_title_tag(markdown_text)
        title = extracted_title or title

        report_data = {
            "results": [{
                "question_id": 1,
                "question": title,
                "title": title,
                "answer": markdown_text,
                "success": True,
                "images": images or [],
                "date_filter": None,
            }],
            "author": author,
            "created_date": created_date,
            "student_id": student_id,
            "department": department,
            "team_name": team_name,
            "role": role,
            "use_template": True,
        }

        if output_path.endswith(".pdf"):
            self.generate_pdf_report(report_data, output_path)
        else:
            self.generate_word_report(report_data, output_path)


def main():
    parser = argparse.ArgumentParser(description="보고서 문서 생성기")
    parser.add_argument("--json", type=str, required=True, help="입력 JSON 파일")
    parser.add_argument("--output", type=str, required=True, help="출력 파일 (.docx 또는 .pdf)")
    args = parser.parse_args()

    with open(args.json, 'r', encoding='utf-8') as f:
        report_data = json.load(f)

    generator = DocumentGenerator()
    if args.output.endswith('.pdf'):
        generator.generate_pdf_report(report_data, args.output)
    elif args.output.endswith('.docx'):
        generator.generate_word_report(report_data, args.output)
    else:
        print("❌ 지원되지 않는 파일 형식입니다. .docx 또는 .pdf를 사용하세요.")


if __name__ == "__main__":
    main()
